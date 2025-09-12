from docx import Document
import os

# Ensure the directory exists
os.makedirs("downloads", exist_ok=True)

# Create a new document
document = Document()
document.add_heading('Mock DOCX Document', 0)

p = document.add_paragraph('This is a mock DOCX file created for testing purposes. ')
p.add_run('It contains some text.').bold = True
p.add_run(' And some ')
p.add_run('italic text.').italic = True

document.add_paragraph('這是用於測試的模擬 DOCX 檔案。')

document.save('downloads/mock_file_2.docx')

print("Mock DOCX file created successfully at downloads/mock_file_2.docx")
