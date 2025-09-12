import pytest
import sys
import sqlite3
import json
import os
import re
from pathlib import Path
from unittest.mock import MagicMock
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
from PIL import Image
import gdown
import filetype

# --- 測試環境路徑設定 ---
SRC_DIR = Path(__file__).resolve().parent.parent / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

# --- 準備匯入被測試的模組 ---
from db.database import get_db_connection, initialize_database
from api.routes.page3_processor import run_processing_task
from api.routes.page4_analyzer import run_ai_analysis_task

# --- Pytest Fixtures (測試輔助工具) ---
# The db_conn fixture is now in conftest.py

@pytest.fixture(scope="module")
def dummy_image_path(tmpdir_factory):
    """建立一個用於測試的假圖片，並回傳其路徑。"""
    img = Image.new('RGB', (100, 100), color = 'blue')
    img_path = tmpdir_factory.mktemp("fixtures").join("dummy_image.png")
    img.save(str(img_path))
    return str(img_path)

def create_test_docx(dir_path, image_path):
    """輔助函式：建立一個測試用的 DOCX 檔案。"""
    doc = Document()
    doc.add_paragraph("這是DOCX文件。")
    doc.add_picture(image_path, width=Inches(1))
    file_path = Path(dir_path) / "test.docx"
    doc.save(file_path)
    return str(file_path)

def create_test_pdf(dir_path, image_path):
    """輔助函式：建立一個測試用的 PDF 檔案。"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    # 使用英文以避免 FPDF 的 Unicode 字型問題，因為本測試的重點是檔案處理而非文字渲染
    pdf.cell(text="This is a PDF document.")
    # 將圖片置於文字下方
    pdf.image(image_path, x=10, y=20, w=50)
    file_path = Path(dir_path) / "test.pdf"
    pdf.output(file_path)
    return str(file_path)

# --- 整合測試案例 ---

@pytest.mark.parametrize(
    "file_creator, file_type",
    [
        (create_test_docx, "DOCX"),
        (create_test_pdf, "PDF"),
    ],
    ids=["process_docx_file", "process_pdf_file"]
)
def test_real_file_processing_task(db_conn, tmp_path, dummy_image_path, file_creator, file_type):
    """
    這是一個參數化的整合測試，用於驗證真實檔案（DOCX 和 PDF）的處理流程。
    """
    # 1. **準備**: 建立測試檔案並在資料庫中設定初始狀態
    # tmp_path 是 pytest 提供的每個測試函式專用的暫存目錄
    test_file_path = file_creator(tmp_path, dummy_image_path)

    cursor = db_conn.cursor()
    # 插入一筆模擬已下載完成的紀錄
    cursor.execute(
        "INSERT INTO extracted_urls (url, status, local_path) VALUES (?, ?, ?)",
        (f"http://test.com/test.{file_type.lower()}", "completed", test_file_path)
    )
    db_conn.commit()
    url_id = cursor.lastrowid
    assert url_id is not None

    # 2. **執行**: 呼叫背景任務函式來處理這個檔案
    # 我們在測試中同步執行它，並傳入一個假的埠號 (port=0)，因為我們不測試通知部分
    run_processing_task(url_id, port=0)

    # 3. **驗證**: 從資料庫中讀取結果並進行斷言
    cursor.execute("SELECT * FROM extracted_urls WHERE id = ?", (url_id,))
    result = cursor.fetchone()

    assert result is not None, "處理後的紀錄不應為 None"

    # 驗證狀態是否已更新
    assert result["status"] == "processed", f"{file_type} 檔案處理後狀態應為 'processed'"

    # 驗證檔案雜湊值是否已計算
    assert result["file_hash"] is not None, f"{file_type} 檔案應計算並儲存 file_hash"
    assert len(result["file_hash"]) == 64, "file_hash 應為 SHA256 的長度"

    # 驗證圖片是否已提取
    assert result["extracted_image_paths"] is not None, f"{file_type} 檔案應提取圖片路徑"

    try:
        extracted_images = json.loads(result["extracted_image_paths"])
        assert isinstance(extracted_images, list), "提取的圖片路徑應為一個列表"
        assert len(extracted_images) > 0, f"應從 {file_type} 檔案中提取出至少一張圖片"

        # 驗證提取出的圖片檔案確實存在
        extracted_image_path = Path(extracted_images[0])
        assert extracted_image_path.exists(), f"提取出的圖片檔案 {extracted_image_path} 應存在於檔案系統中"

    except (json.JSONDecodeError, IndexError):
        pytest.fail("extracted_image_paths 欄位不是一個有效的、包含路徑的 JSON 列表")

    # 驗證文字是否已提取
    assert "extracted_text" in result.keys(), "資料庫紀錄應包含 'extracted_text' 欄位"
    assert result["extracted_text"] is not None, "提取的文字不應為 None"
    assert len(result["extracted_text"]) > 0, "提取的文字內容不應為空"
    if file_type == "DOCX":
        assert "這是DOCX文件" in result["extracted_text"]
    elif file_type == "PDF":
        assert "This is a PDF document" in result["extracted_text"]


    print(f"\n✅ {file_type} 檔案整合測試成功!")
    print(f"  - 狀態更新為: {result['status']}")
    print(f"  - 提取的圖片: {result['extracted_image_paths']}")
    print(f"  - 提取的文字長度: {len(result['extracted_text'])}")


def test_ai_analysis_retry_queue_logic(db_conn, monkeypatch):
    """
    整合測試：驗證 AI 分析的重試佇列邏輯。
    - 第一次分析失敗時，狀態應變為 'pending_retry'。
    - 手動觸發第二次分析時，應能成功並將狀態更新為 'analyzed'。
    """
    # --- 1. 準備 ---

    # 在資料庫中插入一筆可供分析的紀錄
    cursor = db_conn.cursor()
    cursor.execute(
        "INSERT INTO extracted_urls (url, status, local_path, source_text, retry_count) VALUES (?, ?, ?, ?, ?)",
        ("http://test.com/retry_test", "processed", "/fake/path", "這是要分析的文字。", 0)
    )
    db_conn.commit()
    file_id = cursor.lastrowid

    # Mock `prompt_manager` 以避免檔案依賴
    mock_prompt_manager = MagicMock()
    mock_prompt_manager.get_all_prompts.return_value = {
        "stage_1_extraction_prompt": "prompt1 {document_text}",
        "stage_2_generation_prompt": "prompt2 {data_package}"
    }
    monkeypatch.setattr("api.routes.page4_analyzer.prompt_manager", mock_prompt_manager)

    # Mock `key_manager`
    mock_key_manager = MagicMock()
    mock_key_manager.get_all_valid_keys_for_manager.return_value = [{'name': 'mock_key', 'value': '123'}]
    monkeypatch.setattr("api.routes.page4_analyzer.key_manager", mock_key_manager)

    # Mock `GeminiManager`
    mock_gemini_instance = MagicMock()
    # 第一次呼叫 prompt_for_json 時拋出錯誤，第二次正常運作
    mock_gemini_instance.prompt_for_json.side_effect = [
        RuntimeError("模擬 AI 第一次失敗"), # 第一次呼叫失敗
        {"key": "value"} # 第二次呼叫成功
    ]
    mock_gemini_instance.prompt_for_text.return_value = "<html></html>"

    # 將 mock instance 作為類別的回傳值
    monkeypatch.setattr("api.routes.page4_analyzer.GeminiManager", MagicMock(return_value=mock_gemini_instance))

    # --- 2. 執行第一次分析 (預期會失敗並進入重試佇列) ---
    print("\n--- 執行第一次分析 (預期失敗) ---")
    test_model = "test-model-for-retry"
    run_ai_analysis_task([file_id], server_port=0, model_name=test_model)

    # 驗證 prompt_for_json 是用正確的模型名稱呼叫的
    mock_gemini_instance.prompt_for_json.assert_any_call(
        prompt="prompt1 {document_text}".format(document_text="這是要分析的文字。"),
        model_name=test_model
    )

    # --- 3. 驗證第一次分析的結果 ---
    cursor.execute("SELECT status, retry_count, last_error_details FROM extracted_urls WHERE id = ?", (file_id,))
    result1 = cursor.fetchone()

    print(f"第一次分析後狀態: {dict(result1)}")
    assert result1['status'] == 'pending_retry', "第一次失敗後，狀態應為 'pending_retry'"
    assert result1['retry_count'] == 1, "第一次失敗後，重試次數應為 1"
    assert "模擬 AI 第一次失敗" in result1['last_error_details'], "應記錄第一次失敗的錯誤詳情"

    # --- 4. 執行第二次分析 (模擬手動重試，預期成功) ---
    print("\n--- 執行第二次分析 (預期成功) ---")
    run_ai_analysis_task([file_id], server_port=0, model_name=test_model)

    # --- 5. 驗證第二次分析的結果 ---
    cursor.execute("SELECT status, retry_count FROM extracted_urls WHERE id = ?", (file_id,))
    result2 = cursor.fetchone()
    cursor.execute("SELECT id FROM reports WHERE source_url_id = ?", (file_id,))
    report_result = cursor.fetchone()

    print(f"第二次分析後狀態: {dict(result2)}")
    assert result2['status'] == 'analyzed', "第二次重試成功後，狀態應為 'analyzed'"
    # 重試次數不應再增加
    assert result2['retry_count'] == 1, "成功後，重試次數不應再增加"
    assert report_result is not None, "成功後，應在 reports 表中建立一筆紀錄"
    print("✅ AI 分析重試佇列整合測試成功！")


# TODO: 該測試需要更新以匹配 drive_downloader.py 中 download_file 的最新函式簽名。
# 當前的函式需要 author, message_date, 和 message_time，而不是 created_at_str。
# 在 drive_downloader 功能變更後，應取消註解並修復此測試。
#
# def test_drive_downloader_new_logic(tmp_path, monkeypatch):
#     """
#     測試 drive_downloader 是否能根據新邏輯正確處理檔案：
#     - 使用 filetype 偵測副檔名
#     - 使用傳入的 url_id 和 created_at 建立檔名
#     """
#     # 1. 準備
#     fake_download_dir = tmp_path / "downloads"
#     fake_download_dir.mkdir()
#
#     # 模擬 filetype.guess 的回傳物件
#     class MockFileType:
#         extension = "pdf"
#         mime = "application/pdf"
#
#     # 2. Mock: 使用更穩健的方式模擬 gdown 的行為
#     def mock_gdown_download(url, output, **kwargs):
#         # 模擬 gdown 的核心行為：在指定的 output 路徑建立一個非空檔案
#         with open(output, "w") as f:
#             f.write("mock content")
#         # 回傳該路徑，就像真實的 gdown 會做的一樣
#         return output
#
#     monkeypatch.setattr(gdown, "download", mock_gdown_download)
#     monkeypatch.setattr("filetype.guess", lambda path: MockFileType())
#
#     # 3. 執行
#     from tools.drive_downloader import download_file
#     url_id = 99
#     created_at = "2025-01-01 10:30:00"
#
#     final_path_str = download_file(
#         url="http://fake.url/doc.pdf",
#         output_dir=str(fake_download_dir),
#         url_id=url_id,
#         created_at_str=created_at
#     )
#
#     # 4. 驗證
#     assert final_path_str is not None
#     final_path = Path(final_path_str)
#
#     # 驗證最終檔案存在
#     assert final_path.exists()
#
#     # 驗證檔名是否符合 'YYYY-MM-DDTHH-MM-SS_file_ID.ext' 格式
#     expected_timestamp = "2025-01-01T10-30-00"
#     expected_filename = f"{expected_timestamp}_file_{url_id}.pdf"
#     assert final_path.name == expected_filename, f"檔名應為 {expected_filename}，但卻是 {final_path.name}"
