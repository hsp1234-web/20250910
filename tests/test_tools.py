import pytest
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches
from PIL import Image
import os

# --- 測試環境路徑設定 ---
# 確保測試程式可以找到 src 目錄下的模組
# 雖然 pytest.ini 已設定，但在某些執行環境下顯式加入更為保險
SRC_DIR = Path(__file__).resolve().parent.parent / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

# --- 準備匯入被測試的模組 ---
# 這些是我們想要測試的核心工具
from tools.content_extractor import extract_content
from tools.image_compressor import compress_image

# --- Pytest Fixtures (測試輔助工具) ---

@pytest.fixture(scope="module")
def temp_test_dir(tmpdir_factory):
    """
    建立一個專供本模組所有測試使用的暫存目錄。
    'scope="module"' 表示此目錄在所有測試執行前只會建立一次。
    """
    # 使用 pytest 內建的 tmpdir_factory 來建立一個唯一的暫存目錄
    return tmpdir_factory.mktemp("test_tools_data")

@pytest.fixture(scope="module")
def dummy_image_path(temp_test_dir):
    """
    在暫存目錄中建立一個用於測試的假圖片，並回傳其路徑。
    """
    # 建立一個 100x100 的紅色正方形圖片
    img = Image.new('RGB', (100, 100), color = 'red')
    img_path = Path(temp_test_dir) / "dummy_image.png"
    img.save(img_path)
    return str(img_path)

@pytest.fixture(scope="module")
def simulated_docx_path(temp_test_dir, dummy_image_path):
    """
    在暫存目錄中建立一個包含文字和圖片的假 .docx 檔案。
    """
    document = Document()
    document.add_heading('測試文件標題', 0)

    p = document.add_paragraph('這是一個測試段落，用於模擬文件中的文字內容。')
    p.add_run('這是粗體字').bold = True
    p.add_run('，而這是斜體字。').italic = True

    document.add_paragraph('接下來是一張圖片：')
    # 將我們建立的假圖片加入文件中
    document.add_picture(dummy_image_path, width=Inches(1.25))

    docx_path = Path(temp_test_dir) / "simulated_document.docx"
    document.save(str(docx_path))
    return str(docx_path)

# --- 測試案例 ---

def test_file_processing_pipeline_simulation(simulated_docx_path, temp_test_dir):
    """
    這是一個完整的模擬測試案例，用於驗證從內容提取到圖片壓縮的整個流程。
    它使用了上面定義的 fixtures 來取得一個動態生成的 DOCX 檔案。
    """
    # --- 第一階段：驗證內容提取 (extract_content) ---

    # 1. **準備**: 定義提取出的圖片要儲存的目錄
    image_output_dir = Path(temp_test_dir) / "extracted_images"
    image_output_dir.mkdir()

    # 2. **執行**: 呼叫我們要測試的核心函式
    extraction_result = extract_content(simulated_docx_path, str(image_output_dir))

    # 3. **驗證**:
    # 檢查函式是否回傳了包含圖片路徑的字典
    assert extraction_result is not None, "內容提取函式不應回傳 None"
    assert "image_paths" in extraction_result, "提取結果應包含 'image_paths' 鍵"

    image_paths = extraction_result["image_paths"]
    # 檢查是否成功提取了一張圖片
    assert isinstance(image_paths, list), "圖片路徑應為一個列表"
    assert len(image_paths) == 1, "應從文件中提取出一張圖片"

    extracted_image_path_str = image_paths[0]
    extracted_image_path = Path(extracted_image_path_str)

    # 檢查提取出的圖片檔案是否真的存在於檔案系統中
    assert extracted_image_path.exists(), f"提取出的圖片檔案不存在於: {extracted_image_path}"
    # 檢查檔案是否在我們指定的輸出目錄中
    assert image_output_dir in extracted_image_path.parents, "提取出的圖片未儲存在指定的目錄"

    # --- 第二階段：驗證圖片壓縮 (compress_image) ---

    # 1. **準備**: 定義壓縮後圖片的儲存目錄
    compressed_output_dir = Path(temp_test_dir) / "compressed_images"
    compressed_output_dir.mkdir()

    # 2. **執行**: 將第一階段提取出的圖片路徑，傳遞給壓縮函式
    compressed_path_str = compress_image(str(extracted_image_path), str(compressed_output_dir))

    # 3. **驗證**:
    # 檢查函式是否回傳了有效的路徑
    assert compressed_path_str is not None, "圖片壓縮函式不應回傳 None"

    compressed_path = Path(compressed_path_str)
    # 檢查壓縮後的檔案是否真的存在
    assert compressed_path.exists(), f"壓縮後的圖片檔案不存在於: {compressed_path}"
    # 檢查壓縮後的檔案是否在我們指定的輸出目錄中
    assert compressed_output_dir in compressed_path.parents

    # 檢查壓縮是否有效（壓縮後的檔案大小應小於或等於原始檔案）
    original_size = extracted_image_path.stat().st_size
    compressed_size = compressed_path.stat().st_size
    assert compressed_size <= original_size, "壓縮後的圖片大小應小於或等於原始圖片"
    assert compressed_size > 0, "壓縮後的圖片檔案不應為空"

    print(f"\n✅ 模擬測試成功!")
    print(f"  - 原始 DOCX: {simulated_docx_path}")
    print(f"  - 提取的圖片: {extracted_image_path} (大小: {original_size} 位元組)")
    print(f"  - 壓縮後圖片: {compressed_path} (大小: {compressed_size} 位元組)")


from unittest.mock import patch

# --- 測試 drive_downloader ---
# 由於 drive_downloader 會在執行時被匯入，我們需要確保路徑已設定
from tools import drive_downloader

def test_drive_downloader_new_naming_logic(temp_test_dir):
    """
    驗證 drive_downloader.py 中新的、強化的檔案命名邏輯。
    這個測試會模擬下載一個檔案，並檢查最終的檔名是否符合以下格式：
    {created_at_timestamp}_file_{url_id}.{detected_extension}
    """
    # --- 1. 準備 (Arrange) ---

    # 建立一個假的、已下載的檔案，讓 magic 可以偵測它
    # 內容 "Hello" 是一個純文字檔
    fake_downloaded_path = Path(temp_test_dir) / "gdown_temp_file"
    fake_downloaded_path.write_text("Hello")

    # 定義測試用的參數
    test_url = "http://fake.drive.url/123"
    test_output_dir = str(temp_test_dir)
    test_url_id = 42
    # 資料庫傳來的時間戳字串
    test_created_at_str = "2025-09-10 17:30:08"

    # 預期產生的檔名組件
    expected_timestamp = "2025-09-10T17-30-08"
    expected_basename = f"file_{test_url_id}"
    expected_extension = ".pdf" # 我們將模擬 magic 偵測到 PDF
    expected_final_filename = f"{expected_timestamp}_{expected_basename}{expected_extension}"
    expected_final_path = Path(test_output_dir) / expected_final_filename

    # --- 2. 執行 (Act) ---

    # 使用 patch 來模擬外部依賴的行為
    # 我們不希望真的去下載或執行檔案系統操作
    with patch('tools.drive_downloader.gdown.download', return_value=str(fake_downloaded_path)) as mock_gdown, \
         patch('tools.drive_downloader.magic.from_file', return_value='application/pdf') as mock_magic, \
         patch('pathlib.Path.rename') as mock_rename, \
         patch('os.path.exists', return_value=True): # 假設檔案下載後存在

        # 呼叫我們想要測試的函式
        final_path_str = drive_downloader.download_file(
            url=test_url,
            output_dir=test_output_dir,
            url_id=test_url_id,
            created_at=test_created_at_str
        )

    # --- 3. 驗證 (Assert) ---

    # 驗證 gdown.download 是否被正確呼叫
    mock_gdown.assert_called_once_with(test_url, test_output_dir, quiet=False, fuzzy=True)

    # 驗證 magic.from_file 是否被用來偵測我們假的下載檔案
    mock_magic.assert_called_once_with(str(fake_downloaded_path), mime=True)

    # 驗證函式回傳的路徑是否正確
    assert final_path_str is not None, "函式不應回傳 None"
    assert final_path_str == str(expected_final_path), "回傳的最終路徑與預期不符"

    # 驗證 rename 是否被呼叫，且其目標路徑是我們預期的最終路徑
    mock_rename.assert_called_once_with(expected_final_path)

    print(f"\n✅ 檔案命名邏輯測試成功!")
    print(f"  - 輸入 ID: {test_url_id}")
    print(f"  - 輸入時間: '{test_created_at_str}'")
    print(f"  - 模擬偵測類型: 'application/pdf'")
    print(f"  - 預期檔名: '{expected_final_filename}'")
    print(f"  - 函式回傳: '{Path(final_path_str).name}'")
